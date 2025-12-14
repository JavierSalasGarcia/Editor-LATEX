"""
Worker de Windows para procesar documentos
Conexión remota a MySQL de Hostinger
"""

import os
import sys
import io
import time
import tempfile
import subprocess
from datetime import datetime, time as dt_time
from pathlib import Path
import mysql.connector
from mysql.connector import Error
import google.generativeai as genai


# ============= CONFIGURACIÓN =============
# Variables de entorno requeridas
MYSQL_HOST = os.environ.get('MYSQL_HOST', 'localhost')
MYSQL_PORT = int(os.environ.get('MYSQL_PORT', '3306'))
MYSQL_USER = os.environ.get('MYSQL_USER', '')
MYSQL_PASS = os.environ.get('MYSQL_PASS', '')
MYSQL_DB = os.environ.get('MYSQL_DB', 'editor_latex')

GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY', '')

# Carpetas locales
PLANTILLAS_FOLDER = os.environ.get('PLANTILLAS_FOLDER', r'C:\Editor-LATEX\plantillas')
TEMP_FOLDER = os.environ.get('TEMP_FOLDER', r'C:\Editor-LATEX\temp')

# Carpeta de uploads del servidor (puede ser carpeta compartida de red o ruta local)
# Ejemplo: \\servidor\Editor-LATEX\php\uploads (Windows UNC path)
# Ejemplo: /mnt/servidor/Editor-LATEX/php/uploads (Linux mount)
SERVER_UPLOADS_FOLDER = os.environ.get('SERVER_UPLOADS_FOLDER', r'\\localhost\Editor-LATEX\php\uploads')
SERVER_PROCESSED_FOLDER = os.environ.get('SERVER_PROCESSED_FOLDER', r'\\localhost\Editor-LATEX\php\processed')

# Horario de trabajo (8am - 5pm)
WORK_START_HOUR = int(os.environ.get('WORK_START_HOUR', '8'))
WORK_END_HOUR = int(os.environ.get('WORK_END_HOUR', '17'))

# Intervalo de polling (segundos)
POLL_INTERVAL = int(os.environ.get('POLL_INTERVAL', '30'))

# =========================================


class MySQLConnection:
    """Manejo de conexión a MySQL remoto"""

    def __init__(self):
        self.connection = None
        self.connect()

    def connect(self):
        """Establece conexión a MySQL"""
        try:
            self.connection = mysql.connector.connect(
                host=MYSQL_HOST,
                port=MYSQL_PORT,
                user=MYSQL_USER,
                password=MYSQL_PASS,
                database=MYSQL_DB,
                charset='utf8mb4',
                collation='utf8mb4_unicode_ci'
            )
            print(f"✓ Conectado a MySQL: {MYSQL_HOST}:{MYSQL_PORT}")
        except Error as e:
            print(f"✗ Error conectando a MySQL: {e}")
            raise

    def ensure_connected(self):
        """Asegura que la conexión esté activa"""
        try:
            if self.connection and self.connection.is_connected():
                return True
        except:
            pass

        print("⚠ Reconectando a MySQL...")
        self.connect()
        return self.connection.is_connected()

    def get_cursor(self):
        """Obtiene cursor asegurando conexión"""
        self.ensure_connected()
        return self.connection.cursor(dictionary=True)

    def commit(self):
        """Commit de transacción"""
        if self.connection:
            self.connection.commit()

    def close(self):
        """Cierra conexión"""
        if self.connection and self.connection.is_connected():
            self.connection.close()
            print("✓ Conexión MySQL cerrada")


class DocumentProcessor:
    """Procesa documentos usando Gemini API"""

    def __init__(self, api_key, db_connection):
        if not api_key:
            raise ValueError("Se requiere GEMINI_API_KEY")

        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-pro')
        self.db = db_connection

        # Crear carpetas
        os.makedirs(PLANTILLAS_FOLDER, exist_ok=True)
        os.makedirs(TEMP_FOLDER, exist_ok=True)

    def log(self, job_id, level, message):
        """Registra log en base de datos"""
        try:
            cursor = self.db.get_cursor()
            cursor.execute(
                "INSERT INTO processing_logs (job_id, log_level, message) VALUES (%s, %s, %s)",
                (job_id, level, message)
            )
            self.db.commit()
            cursor.close()
        except Exception as e:
            print(f"  ⚠ Error al guardar log: {e}")

    def get_pending_jobs(self):
        """Obtiene trabajos pendientes de la base de datos"""
        try:
            cursor = self.db.get_cursor()
            cursor.execute("""
                SELECT * FROM v_jobs_with_config
                WHERE status = 'pending'
                ORDER BY created_at ASC
                LIMIT 5
            """)
            jobs = cursor.fetchall()
            cursor.close()
            return jobs
        except Exception as e:
            print(f"✗ Error obteniendo trabajos: {e}")
            return []

    def update_job_status(self, job_id, status, error_message=None):
        """Actualiza estado de un trabajo"""
        try:
            cursor = self.db.get_cursor()

            if status == 'processing':
                cursor.execute(
                    "UPDATE jobs SET status = %s, started_at = NOW() WHERE job_id = %s",
                    (status, job_id)
                )
            elif status == 'completed':
                cursor.execute(
                    "UPDATE jobs SET status = %s, completed_at = NOW() WHERE job_id = %s",
                    (status, job_id)
                )
            elif status == 'error':
                cursor.execute(
                    "UPDATE jobs SET status = %s, error_message = %s WHERE job_id = %s",
                    (status, error_message, job_id)
                )

            self.db.commit()
            cursor.close()
        except Exception as e:
            print(f"  ✗ Error actualizando estado: {e}")

    def create_zip_package(self, job_id, pdf_path, temp_dir):
        """
        Crea un archivo ZIP con:
        - Todos los archivos LaTeX (.tex, .cls, logos/, figuras/)
        - PDF compilado
        """
        import zipfile

        print(f"  → Creando paquete ZIP...")

        zip_path = os.path.join(TEMP_FOLDER, f"{job_id}.zip")

        try:
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                # Agregar todos los archivos del directorio temporal
                # Esto incluye: .tex, .cls, logos/, figuras/, etc.
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        # Calcular ruta relativa para mantener estructura
                        arcname = os.path.relpath(file_path, temp_dir)

                        # Solo agregar archivos relevantes (excluir .aux, .log, etc.)
                        if file.endswith(('.tex', '.cls', '.sty', '.bst', '.bib', '.png', '.jpg', '.jpeg', '.pdf')):
                            zipf.write(file_path, arcname)
                            print(f"    Agregando: {arcname}")

            zip_size = os.path.getsize(zip_path)
            print(f"  ✓ ZIP creado ({zip_size / 1024:.2f} KB)")

            return zip_path

        except Exception as e:
            print(f"  ✗ Error creando ZIP: {e}")
            raise

    def upload_zip_to_server(self, job_id, zip_path):
        """Sube ZIP al servidor y actualiza MySQL"""
        try:
            zip_filename = f"{job_id}.zip"
            zip_size = os.path.getsize(zip_path)

            # Copiar ZIP a carpeta processed del servidor
            dest_path = os.path.join(SERVER_PROCESSED_FOLDER, zip_filename)

            # Crear carpeta si no existe
            os.makedirs(SERVER_PROCESSED_FOLDER, exist_ok=True)

            # Copiar archivo
            import shutil
            shutil.copy2(zip_path, dest_path)

            print(f"  ✓ ZIP copiado a servidor: {dest_path}")

            # Actualizar MySQL
            cursor = self.db.get_cursor()
            cursor.execute(
                "UPDATE jobs SET zip_filename = %s, zip_size = %s WHERE job_id = %s",
                (zip_filename, zip_size, job_id)
            )
            self.db.commit()
            cursor.close()

            print(f"  ✓ MySQL actualizado ({zip_size / 1024:.2f} KB)")

        except Exception as e:
            print(f"  ✗ Error subiendo ZIP: {e}")
            raise

    def notify_user(self, job_id, status):
        """Marca trabajo para notificación (PHP enviará el email)"""
        try:
            cursor = self.db.get_cursor()
            cursor.execute(
                "UPDATE jobs SET notified = FALSE WHERE job_id = %s",
                (job_id,)
            )
            self.db.commit()
            cursor.close()
        except Exception as e:
            print(f"  ⚠ Error marcando para notificación: {e}")

    def process_word_to_latex(self, word_data, plantilla_folder, plantilla_main, job):
        """Convierte documento Word a LaTeX usando Gemini"""
        print(f"  → Procesando Word con Gemini...")

        try:
            # Guardar archivo Word temporalmente
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(word_data)
                word_file = tmp.name

            # Leer con python-docx
            try:
                from docx import Document
                doc = Document(word_file)
            except ImportError:
                print("  → Instalando python-docx...")
                subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx'],
                             capture_output=True)
                from docx import Document
                doc = Document(word_file)

            # Extraer texto
            content = '\n'.join([para.text for para in doc.paragraphs])

            # Limpiar archivo temporal
            os.unlink(word_file)

            # Leer plantilla principal (main.tex)
            plantilla_path = os.path.join(PLANTILLAS_FOLDER, plantilla_folder, plantilla_main)
            plantilla_content = ""

            if os.path.exists(plantilla_path):
                with open(plantilla_path, 'r', encoding='utf-8') as f:
                    plantilla_content = f.read()
            else:
                print(f"  ⚠ Plantilla no encontrada: {plantilla_path}")

            # Prompt para Gemini
            prompt = f"""
Eres un experto en LaTeX. Convierte el siguiente documento a formato LaTeX siguiendo esta plantilla.

PLANTILLA DE LA REVISTA:
{plantilla_content or 'Usa una plantilla estándar de artículo académico'}

CONTENIDO DEL DOCUMENTO:
{content}

METADATOS:
- Volumen: {job['volumen']}
- Año: {job['año']}
- Número: {job['numero']}
- Página inicial: {job['pagina_inicial']}

INSTRUCCIONES:
1. Mantén la estructura de la plantilla
2. Inserta el contenido del documento en las secciones apropiadas
3. Convierte formato (negritas, cursivas, listas) a comandos LaTeX
4. Asegúrate de que el código LaTeX sea compilable
5. NO incluyas explicaciones, solo devuelve el código LaTeX

Devuelve SOLO el código LaTeX completo, sin bloques de código markdown.
"""

            print("  → Consultando Gemini API...")
            response = self.model.generate_content(prompt)
            latex_content = response.text

            # Limpiar respuesta
            latex_content = latex_content.replace('```latex', '').replace('```', '').strip()

            return latex_content

        except Exception as e:
            print(f"  ✗ Error procesando Word: {e}")
            raise

    def process_latex_file(self, latex_data, plantilla_folder, plantilla_main, job):
        """Procesa archivo LaTeX existente"""
        print(f"  → Procesando LaTeX...")

        try:
            latex_content = latex_data.decode('utf-8', errors='ignore')

            # Leer plantilla principal
            plantilla_path = os.path.join(PLANTILLAS_FOLDER, plantilla_folder, plantilla_main)
            if os.path.exists(plantilla_path):
                with open(plantilla_path, 'r', encoding='utf-8') as f:
                    plantilla_content = f.read()

                # Ajustar con Gemini
                prompt = f"""
Eres un experto en LaTeX. Ajusta el siguiente documento LaTeX a esta plantilla de revista.

PLANTILLA:
{plantilla_content}

DOCUMENTO ORIGINAL:
{latex_content}

METADATOS:
- Volumen: {job['volumen']}
- Año: {job['año']}
- Número: {job['numero']}
- Página inicial: {job['pagina_inicial']}

INSTRUCCIONES:
1. Adapta el documento a la estructura de la plantilla
2. Mantén todo el contenido original
3. Actualiza los metadatos
4. Asegúrate de que sea compilable

Devuelve SOLO el código LaTeX completo.
"""

                print("  → Ajustando con Gemini API...")
                response = self.model.generate_content(prompt)
                latex_content = response.text.replace('```latex', '').replace('```', '').strip()

            return latex_content

        except Exception as e:
            print(f"  ⚠ Error procesando LaTeX: {e}")
            # Retornar original si falla
            return latex_data.decode('utf-8', errors='ignore')

    def compile_latex(self, latex_content, job_id, compilador, plantilla_folder):
        """Compila LaTeX a PDF"""
        print(f"  → Compilando con {compilador}...")

        # Directorio temporal
        temp_dir = os.path.join(TEMP_FOLDER, job_id)
        os.makedirs(temp_dir, exist_ok=True)

        # IMPORTANTE: Copiar TODA la carpeta de la plantilla al directorio temporal
        # Esto incluye: .cls, logos/, figuras/, y cualquier archivo adicional
        plantilla_source = os.path.join(PLANTILLAS_FOLDER, plantilla_folder)

        if os.path.exists(plantilla_source):
            print(f"  → Copiando plantilla desde {plantilla_source}...")
            import shutil

            # Copiar todos los archivos y subcarpetas
            for item in os.listdir(plantilla_source):
                source = os.path.join(plantilla_source, item)
                dest = os.path.join(temp_dir, item)

                if os.path.isdir(source):
                    # Copiar carpeta completa (logos/, figuras/, etc.)
                    if os.path.exists(dest):
                        shutil.rmtree(dest)
                    shutil.copytree(source, dest)
                else:
                    # Copiar archivo (.cls, .sty, etc.)
                    shutil.copy2(source, dest)

            print(f"  ✓ Plantilla copiada")
        else:
            print(f"  ⚠ Carpeta de plantilla no encontrada: {plantilla_source}")

        # Escribir .tex con el contenido procesado
        tex_file = os.path.join(temp_dir, f"{job_id}.tex")
        with open(tex_file, 'w', encoding='utf-8') as f:
            f.write(latex_content)

        # Compilar (2 pasadas)
        for i in range(2):
            print(f"    Pasada {i+1}/2...")
            result = subprocess.run(
                [compilador, '-interaction=nonstopmode', f"{job_id}.tex"],
                cwd=temp_dir,
                capture_output=True,
                text=True,
                timeout=120
            )

            if result.returncode != 0:
                print(f"  ✗ Error en compilación:")
                print(result.stdout[-1000:])
                raise Exception(f"Error al compilar LaTeX (código {result.returncode})")

        # Verificar PDF
        pdf_file = os.path.join(temp_dir, f"{job_id}.pdf")
        if not os.path.exists(pdf_file):
            raise Exception("El PDF no se generó")

        print(f"  ✓ PDF generado")
        return pdf_file

    def process_job(self, job):
        """Procesa un trabajo completo"""
        job_id = job['job_id']

        print(f"\n{'='*60}")
        print(f"Procesando: {job_id}")
        print(f"Archivo: {job['filename_original']}")
        print(f"Revista: {job['revista_nombre']}")
        print(f"{'='*60}")

        try:
            # Actualizar a processing
            self.update_job_status(job_id, 'processing')
            self.log(job_id, 'info', 'Iniciando procesamiento')

            # Leer archivo del servidor
            upload_path = os.path.join(SERVER_UPLOADS_FOLDER, job['upload_filename'])

            if not os.path.exists(upload_path):
                raise Exception(f"Archivo no encontrado en servidor: {upload_path}")

            with open(upload_path, 'rb') as f:
                file_data = f.read()

            # Procesar según tipo
            if job['file_extension'] in ['doc', 'docx']:
                latex_content = self.process_word_to_latex(
                    file_data,
                    job['plantilla_folder'],
                    job['plantilla_main'],
                    job
                )
            elif job['file_extension'] == 'tex':
                latex_content = self.process_latex_file(
                    file_data,
                    job['plantilla_folder'],
                    job['plantilla_main'],
                    job
                )
            else:
                raise Exception(f"Tipo de archivo no soportado: {job['file_extension']}")

            self.log(job_id, 'info', 'Documento procesado con Gemini')

            # Compilar (pasa plantilla_folder para copiar archivos .cls, logos, etc.)
            pdf_file = self.compile_latex(
                latex_content,
                job_id,
                job['compilador'],
                job['plantilla_folder']
            )
            self.log(job_id, 'info', 'PDF compilado')

            # Directorio temporal donde se compiló
            temp_dir = os.path.join(TEMP_FOLDER, job_id)

            # Crear ZIP con LaTeX + PDF
            zip_path = self.create_zip_package(job_id, pdf_file, temp_dir)
            self.log(job_id, 'info', 'ZIP creado con LaTeX + PDF')

            # Subir ZIP al servidor
            self.upload_zip_to_server(job_id, zip_path)
            self.log(job_id, 'info', 'ZIP subido al servidor')

            # Actualizar estado
            self.update_job_status(job_id, 'completed')
            self.log(job_id, 'info', 'Trabajo completado')

            # Marcar para notificación
            self.notify_user(job_id, 'completed')

            print(f"✓ Completado: {job_id}")

        except Exception as e:
            error_msg = str(e)
            print(f"✗ Error: {error_msg}")

            self.update_job_status(job_id, 'error', error_msg)
            self.log(job_id, 'error', error_msg)
            self.notify_user(job_id, 'error')


def is_work_hours():
    """Verifica si estamos en horario de trabajo"""
    now = datetime.now()
    current_hour = now.hour

    # Verificar que sea día laboral (lunes=0, domingo=6)
    if now.weekday() >= 5:  # Sábado o domingo
        return False

    return WORK_START_HOUR <= current_hour < WORK_END_HOUR


def main():
    """Función principal del worker"""
    print("="*60)
    print("WORKER DE PROCESAMIENTO DE DOCUMENTOS")
    print("="*60)
    print(f"MySQL: {MYSQL_HOST}:{MYSQL_PORT}")
    print(f"Base de datos: {MYSQL_DB}")
    print(f"Horario: {WORK_START_HOUR}:00 - {WORK_END_HOUR}:00")
    print(f"Polling: cada {POLL_INTERVAL} segundos")
    print("="*60)

    # Validar configuración
    if not MYSQL_USER or not MYSQL_PASS:
        print("\n✗ ERROR: Configura MYSQL_USER y MYSQL_PASS")
        sys.exit(1)

    if not GEMINI_API_KEY:
        print("\n✗ ERROR: Configura GEMINI_API_KEY")
        sys.exit(1)

    # Conectar a MySQL
    db = MySQLConnection()

    # Inicializar procesador
    processor = DocumentProcessor(GEMINI_API_KEY, db)

    print("\n✓ Worker iniciado. Esperando trabajos...\n")

    try:
        while True:
            # Verificar horario
            if not is_work_hours():
                now = datetime.now()
                print(f"[{now.strftime('%H:%M:%S')}] Fuera de horario laboral. Esperando...")
                time.sleep(300)  # Esperar 5 minutos
                continue

            # Obtener trabajos pendientes
            jobs = processor.get_pending_jobs()

            if jobs:
                print(f"\n✓ Encontrados {len(jobs)} trabajo(s) pendiente(s)")

                for job in jobs:
                    processor.process_job(job)

                print("\n✓ Todos los trabajos procesados")
            else:
                now = datetime.now()
                print(f"[{now.strftime('%H:%M:%S')}] Sin trabajos pendientes. Esperando...")

            # Esperar antes del próximo ciclo
            time.sleep(POLL_INTERVAL)

    except KeyboardInterrupt:
        print("\n\n✓ Deteniendo worker...")
    except Exception as e:
        print(f"\n\n✗ Error fatal: {e}")
    finally:
        db.close()

    print("Worker detenido.")


if __name__ == '__main__':
    main()
